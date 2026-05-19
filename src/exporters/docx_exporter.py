"""DOCX exporter for MVP analysis reports."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Sequence

from src.exporters.schema import NormalizedExportRow, REPORT_TABLE_COLUMNS


class DocxExporter:
    """Create a new Word report with one 7-column results table."""

    output_format = "docx"

    def export(
        self,
        rows: Sequence[NormalizedExportRow],
        output_file: str | Path,
        *,
        source_file: str | Path | None = None,
        run_id: str = "",
    ) -> Path:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover - dependency is in requirements
            raise RuntimeError(
                "python-docx is required to export DOCX reports. "
                "Install it with `pip install python-docx`."
            ) from exc

        output_path = Path(output_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading(f"Результат анализа ТЗ — {_report_date()}", level=1)
        document.add_paragraph(f"RunID: {run_id}")
        if source_file:
            document.add_paragraph(f"Источник: {source_file}")

        table = document.add_table(rows=1, cols=len(REPORT_TABLE_COLUMNS))
        try:
            table.style = "Table Grid"
        except KeyError:  # pragma: no cover - depends on local Word template
            pass

        for cell, title in zip(table.rows[0].cells, REPORT_TABLE_COLUMNS):
            cell.text = title

        for row in rows:
            cells = table.add_row().cells
            values = [
                str(row.id),
                row.ref,
                row.source_text,
                row.status,
                row.comment,
                f"{row.confidence:.2f}",
                row.run_id,
            ]
            for cell, value in zip(cells, values):
                cell.text = value

        document.save(output_path)
        return output_path


def _report_date() -> str:
    return datetime.now(timezone.utc).date().isoformat()
