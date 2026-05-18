"""DOCX parser for tender requirements (ТЗ).

Extracts atomic requirements from a ``.docx`` file by walking paragraphs and
tables. Each non-empty text block (after trimming) becomes a requirement.
For tables, every cell is treated as an independent candidate.

The parser is intentionally permissive — it is designed for the MVP phase
where ТЗ layouts vary widely. Downstream the LLM classifier is responsible
for handling noisy fragments.

If ``python-docx`` is not installed, the parser raises :class:`ParserError`
so callers can degrade gracefully (e.g. ask the user to upload an ``.xlsx``).
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple, Union

from src.parsers.base_parser import BaseParser, ParserError
from src.parsers.excel_parser import load_config

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parse requirements from a ``.docx`` file into a list of dicts."""

    def __init__(self, config_path: Optional[Union[str, Path]] = None) -> None:
        config = load_config(config_path)
        text_config = config.get("text_processing", {})
        docx_config = config.get("docx_parser", {})
        super().__init__(
            min_length=int(
                docx_config.get(
                    "min_length", text_config.get("min_length", self.DEFAULT_MIN_LENGTH)
                )
            ),
            trim=bool(text_config.get("trim", True)),
        )
        self.inline_markers = list(docx_config.get("inline_markers", []))
        self.list_marker_patterns = [
            re.compile(pattern)
            for pattern in docx_config.get("list_marker_patterns", [])
        ]
        self.service_break_patterns = [
            re.compile(pattern)
            for pattern in docx_config.get(
                "service_break_patterns",
                [r"^\s*$", r"^\f+$", r"^[-–—\s]+$"],
            )
        ]

    def load_requirements(
        self, file_path: Union[str, Path]
    ) -> List[Dict[str, Any]]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        if path.suffix.lower() != ".docx":
            raise ParserError(f"Expected a .docx file, got: {path.suffix}")

        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover - depends on optional dep
            raise ParserError(
                "python-docx is required to parse .docx files. "
                "Install it with `pip install python-docx`."
            ) from exc

        try:
            document = Document(str(path))
        except Exception as exc:  # noqa: BLE001
            raise ParserError(f"Failed to read DOCX file {path}: {exc}") from exc

        candidates: List[Tuple[str, Dict[str, Any]]] = []
        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            locator = {"type": "paragraph", "index": paragraph_index}
            candidates.extend(self._text_candidates(paragraph.text, locator))
        for table_index, table in enumerate(document.tables, start=1):
            candidates.extend(self._table_candidates(table, table_index=table_index))

        requirements: List[Dict[str, Any]] = []
        for text, locator in candidates:
            if not self._keep(text):
                continue
            requirements.append(
                {
                    "id": len(requirements) + 1,
                    "text": text,
                    "locator": locator,
                }
            )

        if not requirements:
            raise ParserError(f"DOCX file {path} did not yield any requirements")

        logger.info("Loaded %d requirements from %s", len(requirements), path)
        return requirements

    def _table_candidates(
        self, table: Any, *, table_index: int
    ) -> Iterable[Tuple[str, Dict[str, Any]]]:
        for row_index, row in enumerate(table.rows, start=1):
            for col_index, cell in enumerate(row.cells, start=1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                    locator = {
                        "type": "table",
                        "table": table_index,
                        "row": row_index,
                        "col": col_index,
                        "paragraph": paragraph_index,
                    }
                    yield from self._text_candidates(paragraph.text, locator)
                for nested_index, nested_table in enumerate(cell.tables, start=1):
                    nested_table_index = int(f"{table_index}{nested_index}")
                    yield from self._table_candidates(
                        nested_table, table_index=nested_table_index
                    )

    def _text_candidates(
        self, raw_text: str, locator: Dict[str, Any]
    ) -> Iterable[Tuple[str, Dict[str, Any]]]:
        fragments = str(raw_text or "").splitlines()
        for fragment_index, fragment in enumerate(fragments, start=1):
            text = self._clean_text(fragment)
            if not text or self._is_service_break(text):
                continue
            current_locator = dict(locator)
            if fragment_index > 1:
                current_locator["fragment"] = fragment_index
            yield text, current_locator

    def _clean_text(self, text: str) -> str:
        cleaned = self._normalize(text)
        for marker in self.inline_markers:
            cleaned = cleaned.replace(str(marker), "")
        for pattern in self.list_marker_patterns:
            cleaned = pattern.sub("", cleaned, count=1)
        return self._normalize(cleaned)

    def _is_service_break(self, text: str) -> bool:
        return any(pattern.search(text) for pattern in self.service_break_patterns)


def load_requirements(
    file_path: Union[str, Path],
) -> List[Dict[str, Any]]:
    """Module-level helper used by the central parsers dispatcher."""
    return DocxParser().load_requirements(file_path)
